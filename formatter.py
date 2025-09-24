from docx import Document

def fsd_to_docx(fsd: dict, filename="FSD.docx"):
    doc = Document()
    doc.add_heading(fsd.get("title", "Functional Spec"), 0)

    doc.add_heading("Objective", level=1)
    doc.add_paragraph(fsd.get("objective", ""))

    doc.add_heading("Module", level=1)
    doc.add_paragraph(fsd.get("module", ""))

    doc.add_heading("Inputs", level=1)
    for i in fsd.get("inputs", []):
        doc.add_paragraph(i, style="List Bullet")

    doc.add_heading("Outputs", level=1)
    for o in fsd.get("outputs", []):
        doc.add_paragraph(o, style="List Bullet")

    doc.add_heading("Business Rules", level=1)
    for br in fsd.get("business_rules", []):
        doc.add_paragraph(br, style="List Bullet")

    doc.add_heading("Database Tables", level=1)
    for t in fsd.get("db_tables", []):
        doc.add_paragraph(t, style="List Bullet")

    doc.add_heading("Error Handling", level=1)
    for e in fsd.get("error_handling", []):
        doc.add_paragraph(e, style="List Bullet")

    doc.save(filename)
    return filename
